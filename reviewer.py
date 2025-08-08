import io
import re
import json
from collections import defaultdict
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

DOC_TYPE_KEYWORDS = {
    "Articles of Association": ["articles of association", "aoa", "objects of the company", "share capital"],
    "Memorandum of Association": ["memorandum of association", "moa", "subscribers"],
    "Incorporation Application": ["incorporation application", "application for incorporation"],
    "UBO Declaration": ["ultimate beneficial owner", "ubo", "beneficial owner"],
    "Register of Members and Directors": ["register of members", "register of directors", "members register"],
}

REQUIRED_FOR_INCORP = [
    "Articles of Association",
    "Memorandum of Association",
    "Incorporation Application",
    "UBO Declaration",
    "Register of Members and Directors",
]

JURISDICTION_PATTERNS = [
    (re.compile(r"\bADGM\b", re.I), None),
    (re.compile(r"\bUAE Federal Court(s)?\b|\bDubai Courts\b|\bAbu Dhabi Courts\b", re.I), "Incorrect jurisdiction reference"),
]

SIGNATURE_PATTERNS = [re.compile(r"\bsignature\b|\bsigned\b|\bfor and on behalf\b", re.I)]
AMBIGUOUS_WORDS = [r"\bmay\b", r"\bshould\b"]

def extract_text(doc: Document):
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append((p, text))
    return paragraphs

def detect_doc_type(text):
    scores = defaultdict(int)
    t = text.lower()
    for typ, kws in DOC_TYPE_KEYWORDS.items():
        for kw in kws:
            if kw in t:
                scores[typ] += 1
    if scores:
        typ, val = max(scores.items(), key=lambda kv: kv[1])
        if val > 0:
            return typ
    return "Unknown"

def find_issues(doc: Document):
    paragraphs = extract_text(doc)
    issues = []
    full_text = "\n".join(t for _, t in paragraphs)
    for pattern, message in JURISDICTION_PATTERNS:
        for p, text in paragraphs:
            if pattern.search(text):
                if message:
                    issues.append({"document":"(current)","section":None,"text":text,"issue":message,"severity":"High","loc_para":p})
    if not any(pat.search(full_text) for pat in SIGNATURE_PATTERNS):
        issues.append({"document":"(current)","section":None,"text":"No signature block detected","issue":"Missing signatory block","severity":"High","loc_para":None})
    for p, text in paragraphs:
        for w in AMBIGUOUS_WORDS:
            if re.search(w, text, re.I):
                issues.append({"document":"(current)", "section":None, "text":text, "issue":"Ambiguous language (use shall instead of may/should)", "severity":"Medium", "loc_para":p})
                break
    return issues

def annotate_docx(original_bytes):
    doc = Document(io.BytesIO(original_bytes))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    detected = detect_doc_type(all_text)
    issues = find_issues(doc)
    comment_list = []
    for i, issue in enumerate(issues, start=1):
        para = issue.get("loc_para")
        note = f"«REVIEWER-{i}: {issue['issue']} (severity: {issue['severity']})»"
        if para is not None:
            for run in para.runs:
                try:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                except Exception:
                    pass
            para.add_run(" " + note)
        else:
            p = doc.add_paragraph()
            p.add_run(note)
        comment_list.append({
            "id": f"REVIEWER-{i}",
            "issue": issue['issue'],
            "severity": issue['severity'],
            "text_excerpt": (issue.get('text') or "")[:300]
        })
    doc.add_page_break()
    doc.add_paragraph("Reviewer Comments Summary").runs[0].bold = True
    if comment_list:
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Comment ID"
        hdr_cells[1].text = "Issue"
        hdr_cells[2].text = "Severity"
        hdr_cells[3].text = "Text Excerpt"
        for c in comment_list:
            row_cells = table.add_row().cells
            row_cells[0].text = c["id"]
            row_cells[1].text = c["issue"]
            row_cells[2].text = c["severity"]
            row_cells[3].text = c["text_excerpt"]
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    report = {
        "document_detected_type": detected,
        "issues_found": [
            {"issue": c["issue"], "severity": c["severity"], "id": c["id"], "excerpt": c["text_excerpt"]} for c in comment_list
        ]
    }
    return out.read(), report

def process_files(files):
    annotated_files = []
    reports = []
    type_counts = defaultdict(int)
    for uploaded_file in files:
        content = uploaded_file.read()
        annotated_bytes, report = annotate_docx(content)
        doc_type = report.get("document_detected_type","Unknown")
        type_counts[doc_type] += 1
        annotated_files.append( (uploaded_file.name.replace(".docx","") + "_reviewed.docx", annotated_bytes) )
        reports.append({"file": uploaded_file.name, "report": report})
    uploaded_types = list(type_counts.keys())
    if any(t in ["Articles of Association","Incorporation Application","UBO Declaration"] for t in uploaded_types):
        process = "Company Incorporation"
        required = REQUIRED_FOR_INCORP
        found = sum(1 for r in required if r in uploaded_types)
        missing = [r for r in required if r not in uploaded_types]
    else:
        process = "Unknown"
        required = []
        found = len(uploaded_types)
        missing = []
    final_summary = {
        "process": process,
        "documents_uploaded": len(files),
        "detected_document_types": uploaded_types,
        "required_documents_for_process": required,
        "documents_found": found,
        "missing_documents": missing,
        "file_reports": reports
    }
    return json.dumps(final_summary, indent=2), annotated_files